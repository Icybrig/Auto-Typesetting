"""
Parser for Word (.docx) documents.

Extracts paragraphs, runs, and footnotes from a .docx file and returns a
Document populated with Word-native style names. Style mapping to InDesign
names happens in a separate mapper step.
"""

from __future__ import annotations

from lxml import etree
import docx
from docx import Document as DocxDocument

from src.models.content import Document, Footnote, Paragraph, Run

# Word XML namespace
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


def _get_footnotes_part(doc: DocxDocument) -> dict[int, list[Run]]:
    """
    Build a mapping of footnote id -> list[Run] from the document's footnotes part.

    Word stores footnotes in a separate XML part (word/footnotes.xml). Each
    <w:footnote w:id="N"> element contains one or more <w:p> paragraphs, each
    of which contains <w:r> runs. We flatten all runs across all paragraphs of
    a single footnote into a single list.

    Word always inserts two "special" footnotes with id -1 (separator) and 0
    (continuation separator); we skip those.

    Returns an empty dict if the document has no footnotes part.
    """
    footnotes: dict[int, list[Run]] = {}

    try:
        footnotes_part = doc.part.footnotes_part
    except AttributeError:
        # No footnotes part — document has no footnotes
        return footnotes

    if footnotes_part is None:
        return footnotes

    root = footnotes_part._element  # lxml Element for <w:footnotes>

    for fn_elem in root.findall(f"{W}footnote"):
        raw_id = fn_elem.get(f"{W}id")
        if raw_id is None:
            continue
        fn_id = int(raw_id)
        # Skip Word's built-in separator footnotes (id -1 and 0)
        if fn_id < 1:
            continue

        runs: list[Run] = []
        for para_elem in fn_elem.findall(f".//{W}p"):
            for run_elem in para_elem.findall(f"{W}r"):
                text_parts = [
                    t.text or ""
                    for t in run_elem.findall(f"{W}t")
                ]
                text = "".join(text_parts)

                # Character style for the footnote run
                char_style: str | None = None
                rpr = run_elem.find(f"{W}rPr")
                if rpr is not None:
                    style_elem = rpr.find(f"{W}rStyle")
                    if style_elem is not None:
                        char_style = style_elem.get(f"{W}val")

                if text:  # skip empty runs (e.g. footnote reference mark run)
                    runs.append(Run(text=text, style=char_style))

        footnotes[fn_id] = runs

    return footnotes


def _extract_paragraph_footnotes(
    para_elem: etree._Element,
    footnote_map: dict[int, list[Run]],
) -> list[Footnote]:
    """
    Find all <w:footnoteReference> elements in a paragraph's XML and return
    the corresponding Footnote objects from footnote_map.
    """
    result: list[Footnote] = []
    for ref_elem in para_elem.findall(f".//{W}footnoteReference"):
        raw_id = ref_elem.get(f"{W}id")
        if raw_id is None:
            continue
        fn_id = int(raw_id)
        if fn_id in footnote_map:
            result.append(Footnote(ref_index=fn_id, runs=footnote_map[fn_id]))
    return result


def parse(path: str) -> Document:
    """
    Parse a .docx file and return a Document with Word-native style names.

    Args:
        path: Filesystem path to the .docx file.

    Returns:
        A Document instance containing all paragraphs, their runs, and any
        attached footnotes.

    Raises:
        FileNotFoundError: If the file at *path* does not exist (propagated
            from python-docx / the OS).
        RuntimeError: If python-docx raises any other exception during
            parsing, re-raised with filename and paragraph index context.
    """
    if not __import__("os").path.exists(path):
        raise FileNotFoundError(f"No such file: {path}")
    doc = docx.Document(path)  # raises PackageNotFoundError / IOError if missing

    # Pre-build the footnote map so we can attach footnotes to paragraphs
    footnote_map = _get_footnotes_part(doc)

    paragraphs: list[Paragraph] = []

    for idx, para in enumerate(doc.paragraphs):
        try:
            # --- paragraph style ---
            style_name: str = para.style.name if para.style else "Normal"

            # --- runs ---
            runs: list[Run] = []
            for run in para.runs:
                char_style: str | None = None
                if run.style and run.style.name:
                    # python-docx exposes the *linked* character style name.
                    # Only record it when it differs from "Default Paragraph Font"
                    # which is Word's way of saying "no explicit character style".
                    if run.style.name != "Default Paragraph Font":
                        char_style = run.style.name
                runs.append(Run(text=run.text, style=char_style))

            # --- footnotes ---
            footnotes = _extract_paragraph_footnotes(para._element, footnote_map)

            paragraphs.append(
                Paragraph(runs=runs, style=style_name, footnotes=footnotes)
            )

        except Exception as exc:
            raise RuntimeError(
                f"Error parsing {path} at paragraph {idx}: {exc}"
            ) from exc

    return Document(paragraphs=paragraphs)
