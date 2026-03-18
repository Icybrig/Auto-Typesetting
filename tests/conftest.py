"""Shared pytest fixtures for the Auto-Typesetting test suite.

All .docx fixture files are created on-the-fly using python-docx so no
binary blobs need to be committed to the repository.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
import docx as python_docx


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    """Create an in-memory .docx file.

    ``paragraphs`` is a list of (text, style_name) tuples.  Each tuple
    produces one paragraph.  Returns raw .docx bytes.
    """
    document = python_docx.Document()
    for text, style in paragraphs:
        document.add_paragraph(text, style=style)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# .docx fixtures written to tmp_path
# ---------------------------------------------------------------------------

@pytest.fixture()
def docx_headings(tmp_path: Path) -> Path:
    """A .docx with one Heading 1 and one Heading 2 paragraph."""
    data = _make_docx([
        ("Chapter One", "Heading 1"),
        ("Section One", "Heading 2"),
    ])
    p = tmp_path / "headings.docx"
    p.write_bytes(data)
    return p


@pytest.fixture()
def docx_normal(tmp_path: Path) -> Path:
    """A .docx with two Normal-style paragraphs."""
    data = _make_docx([
        ("First paragraph.", "Normal"),
        ("Second paragraph.", "Normal"),
    ])
    p = tmp_path / "normal.docx"
    p.write_bytes(data)
    return p


@pytest.fixture()
def docx_runs(tmp_path: Path) -> Path:
    """A .docx where a single paragraph has multiple runs with different styles.

    python-docx does not expose a simple "add run with character style" API
    that works reliably without a pre-existing style in the document, so we
    create two separate paragraphs each containing one run and verify that
    runs in a paragraph are captured correctly.
    """
    document = python_docx.Document()
    para = document.add_paragraph(style="Normal")
    run1 = para.add_run("Hello ")
    run2 = para.add_run("World")
    # Applying bold/italic via run properties rather than named styles avoids
    # needing the style to exist; character style will be None for these runs.
    run1.bold = False
    run2.bold = True
    buf = io.BytesIO()
    document.save(buf)
    p = tmp_path / "runs.docx"
    p.write_bytes(buf.getvalue())
    return p


@pytest.fixture()
def docx_empty_paragraph(tmp_path: Path) -> Path:
    """A .docx with one empty (no text) paragraph."""
    document = python_docx.Document()
    document.add_paragraph("", style="Normal")
    buf = io.BytesIO()
    document.save(buf)
    p = tmp_path / "empty_para.docx"
    p.write_bytes(buf.getvalue())
    return p


# ---------------------------------------------------------------------------
# Minimal .indt template fixture (a ZIP file)
# ---------------------------------------------------------------------------

DESIGNMAP_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Document xmlns:idPkg="http://ns.adobe.com/AdobeInDesign/idml/1.0/packaging"
          DOMVersion="16.0">
    <idPkg:Story src="Stories/Story_test.xml"/>
</Document>
""".encode("utf-8")

STORY_XML = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<idPkg:Story xmlns:idPkg="http://ns.adobe.com/AdobeInDesign/idml/1.0/packaging"
             DOMVersion="16.0">
    <Story Self="test">
        <ParagraphStyleRange AppliedParagraphStyle="ParagraphStyle/Body Text">
            <CharacterStyleRange AppliedCharacterStyle="CharacterStyle/$ID/[No character style]">
                <Content>Placeholder</Content>
            </CharacterStyleRange>
        </ParagraphStyleRange>
    </Story>
</idPkg:Story>
""".encode("utf-8")


@pytest.fixture()
def indt_template(tmp_path: Path) -> Path:
    """A minimal .indt file (ZIP) containing designmap.xml and a story."""
    p = tmp_path / "template.indt"
    with zipfile.ZipFile(p, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("designmap.xml", DESIGNMAP_XML)
        zf.writestr("Stories/Story_test.xml", STORY_XML)
    return p


@pytest.fixture()
def not_a_zip(tmp_path: Path) -> Path:
    """A file that is not a valid ZIP archive."""
    p = tmp_path / "notazip.indt"
    p.write_bytes(b"this is not a zip file at all")
    return p


# ---------------------------------------------------------------------------
# Config path
# ---------------------------------------------------------------------------

@pytest.fixture()
def style_map_path() -> str:
    """Absolute path to the project's style_map.json."""
    return str(Path(__file__).parent.parent / "config" / "style_map.json")
